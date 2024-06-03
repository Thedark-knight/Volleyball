from kivy.app import App
from kivy.uix.label import Label
from kivy.uix.boxlayout import BoxLayout
from kivy.uix.button import Button
from kivy.uix.textinput import TextInput
from kivy.uix.filechooser import FileChooserIconView
from kivy.uix.popup import Popup
from kivy.uix.progressbar import ProgressBar
from kivy.clock import Clock
from PIL import Image, ImageDraw, ImageFont
import os
import zipfile
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

class VolleyballApp(App):
    def build(self):
        self.layout = BoxLayout(orientation='vertical')
        self.label = Label(text='Volleyball Rotation Document Generator')
        self.layout.add_widget(self.label)
        
        self.image_button = Button(text='Select Court Image')
        self.image_button.bind(on_press=self.select_image)
        self.layout.add_widget(self.image_button)

        self.title_input = TextInput(hint_text='Enter Document Title')
        self.layout.add_widget(self.title_input)
        
        self.save_button = Button(text='Select Save Directory')
        self.save_button.bind(on_press=self.select_directory)
        self.layout.add_widget(self.save_button)

        self.player_inputs = []
        positions = ["7", "8", "9", "1", "6", "5", "4", "3", "2"]
        for pos in positions:
            input_box = TextInput(hint_text=f'Player for position {pos}')
            self.layout.add_widget(input_box)
            self.player_inputs.append(input_box)
        
        self.generate_button = Button(text='Generate Document')
        self.generate_button.bind(on_press=self.generate_document)
        self.layout.add_widget(self.generate_button)

        self.progress_bar = ProgressBar(max=10)
        self.layout.add_widget(self.progress_bar)
        
        return self.layout

    def select_image(self, instance):
        content = FileChooserIconView()
        popup = Popup(title='Select Court Image', content=content, size_hint=(0.9, 0.9))
        content.bind(on_submit=lambda instance, selection, touch: self.on_image_selected(popup, selection))
        popup.open()

    def on_image_selected(self, popup, selection):
        if selection:
            self.court_image_path = selection[0]
        popup.dismiss()

    def select_directory(self, instance):
        content = FileChooserIconView(dirselect=True)
        popup = Popup(title='Select Save Directory', content=content, size_hint=(0.9, 0.9))
        content.bind(on_submit=lambda instance, selection, touch: self.on_directory_selected(popup, selection))
        popup.open()

    def on_directory_selected(self, popup, selection):
        if selection:
            self.save_directory = selection[0]
        popup.dismiss()

    def generate_document(self, instance):
        players = [input_box.text for input_box in self.player_inputs]
        if not all(players):
            self.label.text = "All player positions must be filled."
            return

        document_title = self.title_input.text
        if not document_title:
            self.label.text = "Document title cannot be empty."
            return

        rotation_order = [7, 8, 9, 1, 6, 5, 4, 3, 2]
        rotations = self.generate_rotations(players, rotation_order, num_rotations=10)

        output_dir = os.path.join(self.save_directory, "volleyball_rotations")
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        output_zip_path = os.path.join(self.save_directory, "volleyball_rotations.zip")

        court_image = Image.open(self.court_image_path)
        font = ImageFont.load_default()

        positions = {
            "1": (300, 200),
            "2": (300, 50),
            "3": (200, 50),
            "4": (100, 50),
            "5": (100, 200),
            "6": (200, 175),
            "7": (350, 130),
            "8": (350, 140),
            "9": (350, 150)
        }

        self.progress_bar.value = 0
        Clock.schedule_interval(lambda dt: self.update_progress(rotations, court_image, font, positions, output_dir, output_zip_path, document_title), 0.1)

    def update_progress(self, rotations, court_image, font, positions, output_dir, output_zip_path, document_title, dt):
        if self.progress_bar.value >= 10:
            self.progress_bar.value = 0
            self.create_word_doc(output_zip_path, document_title)
            return False

        i = int(self.progress_bar.value)
        img = court_image.copy()
        draw = ImageDraw.Draw(img)

        for pos, player in rotations[i].items():
            x, y = positions[pos]
            draw.text((x, y), player, fill="black", font=font)

        img_path = f"{output_dir}/Rotation_{i+1}.png"
        img.save(img_path)

        self.progress_bar.value += 1
        return True

    def generate_rotations(self, players, rotation_order, num_rotations=10):
        rotations = []
        num_positions = len(rotation_order)
        for i in range(num_rotations):
            rotation = {}
            for j in range(num_positions):
                position_index = (i + j) % num_positions
                position = str(rotation_order[position_index])
                rotation[position] = players[j]
            rotations.append(rotation)
        return rotations

    def create_word_doc(self, output_zip_path, title):
        with zipfile.ZipFile(output_zip_path, 'w') as zipf:
            for root, _, files in os.walk(self.save_directory):
                for file in files:
                    zipf.write(os.path.join(root, file), arcname=file)

        extract_to = os.path.join(self.save_directory, 'extracted_images')
        os.makedirs(extract_to, exist_ok=True)
        image_paths = unzip_images(output_zip_path, extract_to)
        new_document = os.path.join(self.save_directory, title + ".docx")
        create_word_doc_with_images(image_paths, new_document, title)
        self.label.text = f"Document created successfully: {new_document}"

def unzip_images(zip_file_path, extract_to):
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        zip_ref.extractall(extract_to)
    return [os.path.join(extract_to, name) for name in zip_ref.namelist()]

def create_word_doc_with_images(image_paths, output_path, title):
    doc = Document()
    doc.add_heading(title, level=1)

    pattern = re.compile(r'Rotation_(\d+)\.png')
    filtered_image_paths = []
    for path in image_paths:
        match = pattern.search(os.path.basename(path))
        if match:
            filtered_image_paths.append((path, int(match.group(1))))

    filtered_image_paths.sort(key=lambda x: x[1])
    sorted_image_paths = [path for path, _ in filtered_image_paths]

    for i in range(0, len(sorted_image_paths), 2):
        image_size = Inches(4.5) if i >= 2 else Inches(4.0)

        if i < len(sorted_image_paths):
            paragraph = doc.add_paragraph(f"Rotation {i + 1}", style='Heading2')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(sorted_image_paths[i], width=image_size)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i + 1 < len(sorted_image_paths):
            paragraph = doc.add_paragraph(f"Rotation {i + 2}", style='Heading2')
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(sorted_image_paths[i + 1], width=image_size)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if i + 2 < len(sorted_image_paths):
            doc.add_page_break()

    doc.save(output_path)

if __name__ == '__main__':
    VolleyballApp().run()
